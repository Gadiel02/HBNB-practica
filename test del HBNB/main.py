from docx import Document
from docx.shared import Inches

def create_technical_document():
    # Crear un nuevo documento de Word
    doc = Document()
    
    # Título
    doc.add_heading('Comprehensive Technical Document for HBnB Project', level=1)

    # Introducción
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(
        "The HBnB project aims to provide a robust platform for users to manage "
        "properties, book places, and share reviews. This technical document serves as "
        "a detailed blueprint for the implementation phases, providing a clear reference "
        "for the system’s architecture and design."
    )

    # High-Level Architecture
    doc.add_heading('2. High-Level Architecture', level=2)
    doc.add_paragraph(
        "The following diagram illustrates the high-level package architecture of the HBnB application, "
        "showing the three main layers: Presentation Layer, Business Logic Layer, and Persistence Layer. "
        "The facade pattern is employed to facilitate communication between these layers."
    )
    
    # Insert High-Level Package Diagram
    doc.add_picture('diagrams/package_diagram.png', width=Inches(6.0))
    doc.add_paragraph(
        "Figure 1: High-Level Package Diagram\n"
        "This diagram presents the layers of the application and their interactions."
    )

    # Business Logic Layer
    doc.add_heading('3. Business Logic Layer', level=2)
    doc.add_paragraph(
        "The detailed class diagram for the Business Logic layer represents the core entities: User, Place, "
        "Review, and Amenity. It illustrates their attributes, methods, and relationships, emphasizing how "
        "they contribute to the application’s functionality."
    )
    
    # Insert Class Diagram
    doc.add_picture('diagrams/class_diagram.png', width=Inches(6.0))
    doc.add_paragraph(
        "Figure 2: Detailed Class Diagram\n"
        "This diagram outlines the entities involved in the business logic layer and their interactions."
    )

    # API Interaction Flow
    doc.add_heading('4. API Interaction Flow', level=2)
    doc.add_paragraph(
        "The following sequence diagrams depict the interactions for key API calls: User Registration, Place Creation, "
        "Review Submission, and Fetching a List of Places. Each diagram illustrates the flow of information and the "
        "sequence of operations required to process user requests."
    )
    
    # Insert Sequence Diagrams
    sequence_diagrams = [
        ('diagrams/user_registration_sequence.png', 'User Registration'),
        ('diagrams/place_creation_sequence.png', 'Place Creation'),
        ('diagrams/review_submission_sequence.png', 'Review Submission'),
        ('diagrams/fetching_places_sequence.png', 'Fetching List of Places')
    ]

    for seq_img, title in sequence_diagrams:
        doc.add_heading(title, level=3)
        doc.add_picture(seq_img, width=Inches(6.0))
        doc.add_paragraph(
            f"Figure X: Sequence Diagram for {title}\n"
            "This diagram outlines the interaction flow for the specified API call, detailing how "
            "the layers communicate to fulfill the request."
        )

    # Guardar el documento
    doc.save('HBnB_Technical_Documentation.docx')
    print("Document created successfully!")

# Ejecutar la función
if __name__ == "__main__":
    create_technical_document()
